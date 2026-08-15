"""
tests/test_formatter.py — Unit tests for thesis_formatter.

Tests build minimal DOCX fixtures in-memory, run the formatter,
and assert output properties without touching the filesystem
(except for a temporary directory for save/reload).
"""
from __future__ import annotations

import io
import os
import sys
import tempfile
import unittest

# Allow running from project root
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from thesis_formatter.analyzer import analyse, ParaType
from thesis_formatter.formatter import format_document
from thesis_formatter import config


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _save_and_reload(doc: Document) -> tuple[str, Document]:
    """Save doc to a temp file and reload. Returns (path, doc)."""
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()
    doc.save(tmp.name)
    return tmp.name, Document(tmp.name)


def _make_doc_with(*paragraphs: tuple[str, str]) -> Document:
    """Create a Document with (text, style) tuples."""
    doc = Document()
    for text, style in paragraphs:
        try:
            p = doc.add_paragraph(text, style=style)
        except Exception:
            p = doc.add_paragraph(text)
    return doc


# ---------------------------------------------------------------------------
# 1. Normal paragraph
# ---------------------------------------------------------------------------
class TestNormalParagraph(unittest.TestCase):
    def test_body_spacing(self):
        doc = _make_doc_with(("Hello world", "Normal"))
        path, _ = _save_and_reload(doc)
        out = path.replace(".docx", "_fmt.docx")
        report = format_document(path, out)
        result = Document(out)
        para = result.paragraphs[0]
        pf   = para.paragraph_format
        self.assertEqual(para.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
        os.unlink(path); os.unlink(out)


# ---------------------------------------------------------------------------
# 2. Bold / Italic paragraph — formatting tests
# ---------------------------------------------------------------------------
class TestBoldItalic(unittest.TestCase):
    def test_heading_bold_preserved(self):
        doc  = Document()
        para = doc.add_paragraph()
        run  = para.add_run("Chapter - I Introduction")
        run.bold = True
        path, _ = _save_and_reload(doc)
        out = path.replace(".docx", "_fmt.docx")
        format_document(path, out)
        result = Document(out)
        bold_runs = [r for p in result.paragraphs
                     for r in p.runs if r.bold]
        self.assertGreater(len(bold_runs), 0, "Heading bold was destroyed!")
        os.unlink(path); os.unlink(out)

    def test_body_bold_stripped_by_default(self):
        doc  = Document()
        para = doc.add_paragraph("This is a normal body paragraph with some ")
        run  = para.add_run("stray bold text")
        run.bold = True
        path, _ = _save_and_reload(doc)
        out = path.replace(".docx", "_fmt.docx")
        format_document(path, out)
        result = Document(out)
        body_runs = [r for p in result.paragraphs
                     for r in p.runs if r.bold]
        self.assertEqual(len(body_runs), 0, "Stray bold in body text was not stripped!")
        os.unlink(path); os.unlink(out)

    def test_shading_and_highlight_stripped(self):
        doc  = Document()
        para = doc.add_paragraph()
        run  = para.add_run("Highlighted copy paste text")
        run.font.highlight_color = 1  # Yellow
        path, _ = _save_and_reload(doc)
        out = path.replace(".docx", "_fmt.docx")
        format_document(path, out)
        result = Document(out)
        for p in result.paragraphs:
            for r in p.runs:
                self.assertIsNone(r.font.highlight_color)
        os.unlink(path); os.unlink(out)

    def test_heading_shading_stripped(self):
        doc  = Document()
        para = doc.add_paragraph("5.4.5 Available nutrient(kg/ha)")
        run  = para.runs[0]
        run.font.highlight_color = 1  # Yellow highlight on heading
        path, _ = _save_and_reload(doc)
        out = path.replace(".docx", "_fmt.docx")
        format_document(path, out)
        result = Document(out)
        for p in result.paragraphs:
            for r in p.runs:
                self.assertIsNone(r.font.highlight_color)
        os.unlink(path); os.unlink(out)

    def test_font_color_reset(self):
        doc  = Document()
        para = doc.add_paragraph("Discolored grey text")
        run  = para.runs[0]
        run.font.color.rgb = RGBColor(89, 89, 89)   # #595959 grey text
        path, _ = _save_and_reload(doc)
        out = path.replace(".docx", "_fmt.docx")
        format_document(path, out)
        result = Document(out)
        para = result.paragraphs[0]
        for r in para.runs:
            # Custom grey color tag should be removed (None / default black)
            self.assertIsNone(r.font.color.rgb)
        os.unlink(path); os.unlink(out)

    def test_complex_bold_stripped(self):
        doc  = Document()
        para = doc.add_paragraph("Normal body text ")
        run  = para.add_run("with complex bold")
        rPr  = run._r.get_or_add_rPr()
        rPr.append(OxmlElement("w:bCs"))
        path, _ = _save_and_reload(doc)
        out = path.replace(".docx", "_fmt.docx")
        format_document(path, out)
        result = Document(out)
        body_runs = [r for p in result.paragraphs for r in p.runs if r.bold]
        self.assertEqual(len(body_runs), 0, "Complex script bold was not stripped!")
        os.unlink(path); os.unlink(out)

    def test_cross_run_and_non_breaking_spaces(self):
        doc = Document()
        para = doc.add_paragraph("The nutrient uptake was increased under elevated CO₂ treatments, \xa0")
        para.add_run(" with the maximum uptake being recorded under T5 (550 ppm CO₂ +1°C) treatment in open field.")
        path, _ = _save_and_reload(doc)
        out = path.replace(".docx", "_fmt.docx")
        format_document(path, out)
        result = Document(out)
        text = result.paragraphs[0].text
        self.assertNotIn("  ", text, "Double space found in formatted paragraph!")
        self.assertNotIn("\xa0", text, "Non-breaking space found in formatted paragraph!")
        os.unlink(path); os.unlink(out)

    def test_italic_preserved(self):
        doc  = Document()
        para = doc.add_paragraph()
        run  = para.add_run("Italic scientific name")
        run.italic = True
        path, _ = _save_and_reload(doc)
        out = path.replace(".docx", "_fmt.docx")
        format_document(path, out)
        result = Document(out)
        italic_runs = [r for p in result.paragraphs
                       for r in p.runs if r.italic]
        self.assertGreater(len(italic_runs), 0, "Italic run was destroyed!")
        os.unlink(path); os.unlink(out)

    def test_stream_processing(self):
        doc = Document()
        para = doc.add_paragraph("This paragraph has double  spaces.")
        in_stream = io.BytesIO()
        doc.save(in_stream)
        in_stream.seek(0)
        
        out_stream = io.BytesIO()
        report = format_document(in_stream, out_stream)
        out_stream.seek(0)
        
        result = Document(out_stream)
        self.assertEqual(len(result.paragraphs), 1)
        self.assertNotIn("  ", result.paragraphs[0].text)
        self.assertEqual(report.para_count, 1)


# ---------------------------------------------------------------------------
# 3. Superscript / Subscript
# ---------------------------------------------------------------------------
class TestSuperSubscript(unittest.TestCase):
    def test_superscript_preserved(self):
        doc  = Document()
        para = doc.add_paragraph("kg ha")
        run  = para.add_run("⁻¹")
        run.font.superscript = True
        path, _ = _save_and_reload(doc)
        out = path.replace(".docx", "_fmt.docx")
        format_document(path, out)
        result = Document(out)
        sup_runs = [r for p in result.paragraphs
                    for r in p.runs if r.font.superscript]
        self.assertGreater(len(sup_runs), 0, "Superscript destroyed!")
        os.unlink(path); os.unlink(out)


# ---------------------------------------------------------------------------
# 4. Heading
# ---------------------------------------------------------------------------
class TestHeading(unittest.TestCase):
    def test_heading1_classified(self):
        doc = _make_doc_with(("4.1 Growth parameters", "Normal"))
        analysis = analyse(doc)
        _, ptype = analysis.classified[0]
        self.assertEqual(ptype, ParaType.HEADING1)

    def test_chapter_classified(self):
        doc = _make_doc_with(("Chapter - IV", "Normal"))
        analysis = analyse(doc)
        _, ptype = analysis.classified[0]
        self.assertEqual(ptype, ParaType.CHAPTER)


# ---------------------------------------------------------------------------
# 5. Bullet list
# ---------------------------------------------------------------------------
class TestBulletList(unittest.TestCase):
    def test_bullet_classified(self):
        doc = _make_doc_with(("- First item", "Normal"))
        analysis = analyse(doc)
        _, ptype = analysis.classified[0]
        self.assertEqual(ptype, ParaType.BULLET)


# ---------------------------------------------------------------------------
# 6. Numbered list
# ---------------------------------------------------------------------------
class TestNumberedList(unittest.TestCase):
    def test_number_classified(self):
        doc = _make_doc_with(("1. First item", "Normal"))
        analysis = analyse(doc)
        _, ptype = analysis.classified[0]
        self.assertEqual(ptype, ParaType.NUMBERED)


# ---------------------------------------------------------------------------
# 7. Table
# ---------------------------------------------------------------------------
class TestTable(unittest.TestCase):
    def test_table_count_preserved(self):
        doc = Document()
        doc.add_table(rows=2, cols=3)
        path, _ = _save_and_reload(doc)
        out = path.replace(".docx", "_fmt.docx")
        format_document(path, out)
        result = Document(out)
        self.assertEqual(len(result.tables), 1)
        os.unlink(path); os.unlink(out)


# ---------------------------------------------------------------------------
# 8. Multi-row table
# ---------------------------------------------------------------------------
class TestMultiRowTable(unittest.TestCase):
    def test_multirow_table_preserved(self):
        doc   = Document()
        table = doc.add_table(rows=5, cols=4)
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.text = f"R{i}C{j}"
        path, _ = _save_and_reload(doc)
        out = path.replace(".docx", "_fmt.docx")
        format_document(path, out)
        result = Document(out)
        self.assertEqual(len(result.tables), 1)
        self.assertEqual(result.tables[0].rows[2].cells[2].text, "R2C2")
        os.unlink(path); os.unlink(out)


# ---------------------------------------------------------------------------
# 9. Image paragraph (drawing XML stub)
# ---------------------------------------------------------------------------
class TestImageParagraph(unittest.TestCase):
    def test_image_classified(self):
        doc  = Document()
        para = doc.add_paragraph()
        # Inject a minimal w:drawing stub
        drawing = OxmlElement("w:drawing")
        inline  = OxmlElement("wp:inline")
        extent  = OxmlElement("wp:extent")
        extent.set("cx", str(int(Inches(8))))   # wide — should be resized
        extent.set("cy", str(int(Inches(4))))
        inline.append(extent)
        drawing.append(inline)
        para._p.append(drawing)

        analysis = analyse(doc)
        _, ptype = analysis.classified[0]
        self.assertEqual(ptype, ParaType.IMAGE)


# ---------------------------------------------------------------------------
# 10. Image + caption
# ---------------------------------------------------------------------------
class TestImageCaption(unittest.TestCase):
    def test_figure_caption_classified(self):
        doc = _make_doc_with(("Figure 4.1.1 Effect of treatment", "Normal"))
        analysis = analyse(doc)
        _, ptype = analysis.classified[0]
        self.assertEqual(ptype, ParaType.FIGURE_CAPTION)


# ---------------------------------------------------------------------------
# 11. Table caption
# ---------------------------------------------------------------------------
class TestTableCaption(unittest.TestCase):
    def test_table_caption_classified(self):
        doc = _make_doc_with(("Table 4.1 Mean values of growth", "Normal"))
        analysis = analyse(doc)
        _, ptype = analysis.classified[0]
        self.assertEqual(ptype, ParaType.TABLE_CAPTION)


# ---------------------------------------------------------------------------
# 12. Page break (blank para removal is limited)
# ---------------------------------------------------------------------------
class TestBlankParas(unittest.TestCase):
    def test_excess_blanks_removed(self):
        doc = Document()
        doc.add_paragraph("Before")
        for _ in range(5):
            doc.add_paragraph("")   # 5 consecutive blanks
        doc.add_paragraph("After")
        path, _ = _save_and_reload(doc)
        out = path.replace(".docx", "_fmt.docx")
        report = format_document(path, out)
        self.assertGreater(report.blanks_removed, 0)
        os.unlink(path); os.unlink(out)


# ---------------------------------------------------------------------------
# 13. Header/footer preserved
# ---------------------------------------------------------------------------
class TestHeaderFooter(unittest.TestCase):
    def test_header_preserved(self):
        doc = Document()
        section = doc.sections[0]
        header  = section.header
        header.paragraphs[0].text = "My Thesis Header"
        path, _ = _save_and_reload(doc)
        out = path.replace(".docx", "_fmt.docx")
        format_document(path, out)
        result = Document(out)
        hdr_text = result.sections[0].header.paragraphs[0].text
        self.assertEqual(hdr_text, "My Thesis Header")
        os.unlink(path); os.unlink(out)


# ---------------------------------------------------------------------------
# 14. Page number field — not converted to static text
# ---------------------------------------------------------------------------
class TestPageNumberField(unittest.TestCase):
    def test_field_preserved(self):
        doc  = Document()
        para = doc.add_paragraph()
        fld  = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        para._p.append(fld)
        instr = OxmlElement("w:instrText")
        instr.text = " PAGE "
        para._p.append(instr)
        fld2 = OxmlElement("w:fldChar")
        fld2.set(qn("w:fldCharType"), "end")
        para._p.append(fld2)

        path, _ = _save_and_reload(doc)
        out = path.replace(".docx", "_fmt.docx")
        format_document(path, out)
        result = Document(out)
        # Check field still exists in XML
        body_xml = result.element.body.xml
        self.assertIn("PAGE", body_xml)
        os.unlink(path); os.unlink(out)


# ---------------------------------------------------------------------------
# 15. Bibliography paragraph
# ---------------------------------------------------------------------------
class TestBibliography(unittest.TestCase):
    def test_bib_classified(self):
        doc = Document()
        doc.add_paragraph("References")
        doc.add_paragraph("Smith, J. (2020). Some paper. Journal, 10(2), 1-10.")
        analysis = analyse(doc)
        types = [pt for _, pt in analysis.classified]
        self.assertIn(ParaType.BIBLIOGRAPHY, types)


# ---------------------------------------------------------------------------
# 16. Front matter
# ---------------------------------------------------------------------------
class TestFrontMatter(unittest.TestCase):
    def test_abstract_classified(self):
        doc = _make_doc_with(("abstract", "Normal"))
        analysis = analyse(doc)
        _, ptype = analysis.classified[0]
        self.assertEqual(ptype, ParaType.FRONT_MATTER)


# ---------------------------------------------------------------------------
if __name__ == "__main__":
    unittest.main(verbosity=2)
