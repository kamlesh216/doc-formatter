import os
import re
import sys
from collections import Counter

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

MARGIN = Inches(1)
LINE_SPACING = 1.15
SPACE_AFTER = Pt(6)
HEADING_COLORS = {
    "Heading 1": "1F4E79",
    "Heading 2": "2E74B5",
    "Heading 3": "2E74B5",
}
HEADING_BUMP = {"Heading 1": 2, "Heading 2": 1, "Heading 3": 0}
LIST_LEVELS = frozenset(["List Bullet", "List Number"])

BULLET_RE = re.compile(r"^\s*[-*•]\s+")
NUMBER_RE = re.compile(r"^\s*(\d+)[.)]\s+")
DOUBLE_SPACE_RE = re.compile(r"\s{2,}")
BLANK_PARA_RE = re.compile(r"^\s*$")


def detect_body_style(doc):
    normal = doc.styles["Normal"]
    body_name = normal.font.name or "Calibri"
    body_size = normal.font.size or Pt(11)
    if not normal.font.name:
        names = Counter()
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name:
                    names[run.font.name] += 1
        if names:
            body_name = names.most_common(1)[0][0]
    return body_name, body_size


def configure_styles(doc, body_name, body_size):
    styles = doc.styles
    for style in styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        if style.name not in HEADING_COLORS:
            continue
        existing = style.font
        if existing.size is None:
            existing.size = Pt(body_size.pt + HEADING_BUMP[style.name])
        if existing.bold is None:
            existing.bold = True
        if existing.color.rgb is None:
            existing.color.rgb = RGBColor.from_string(HEADING_COLORS[style.name])
        if existing.name is None:
            existing.name = body_name

    normal = styles["Normal"]
    if normal.font.name is None:
        normal.font.name = body_name
    if normal.font.size is None:
        normal.font.size = body_size
    pf = normal.paragraph_format
    if pf.line_spacing is None:
        pf.line_spacing = LINE_SPACING
    if pf.space_after is None:
        pf.space_after = SPACE_AFTER


def clean_text(text):
    text = text.replace("\t", " ")
    text = DOUBLE_SPACE_RE.sub(" ", text)
    return text.strip()


def set_alignment(paragraph, style_name):
    if style_name in HEADING_COLORS or style_name in LIST_LEVELS:
        return
    if paragraph.alignment is None:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def paragraph_has_image(paragraph):
    return bool(
        paragraph._p.findall(".//" + qn("w:drawing"))
        or paragraph._p.findall(".//" + qn("w:pict"))
        or paragraph._p.findall(".//" + qn("w:object"))
    )


def make_list(paragraph, regex, new_style):
    text = regex.sub("", paragraph.text).strip()
    paragraph.clear()
    paragraph.add_run(clean_text(text))
    paragraph.style = new_style


def center_image_paragraph(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def format_tables(doc):
    for table in doc.tables:
        table.autofit = True
        current = table.style.name if table.style else None
        if current in (None, "Normal Table", "Table Normal"):
            if "Table Grid" in doc.styles:
                table.style = "Table Grid"
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        cleaned = clean_text(para.text)
                        para.clear()
                        para.add_run(cleaned)


def format_document(input_path, output_path):  # noqa: C901
    if not os.path.isfile(input_path):
        raise FileNotFoundError(f"Input file not found: {input_path}")
    doc = Document(input_path)
    body_name, body_size = detect_body_style(doc)
    configure_styles(doc, body_name, body_size)

    for section in doc.sections:
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN

    previous_was_blank = False
    removed_blank_count = 0
    image_count = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text
        style_name = paragraph.style.name

        if paragraph_has_image(paragraph):
            center_image_paragraph(paragraph)
            image_count += 1
            previous_was_blank = False
            continue

        if BLANK_PARA_RE.match(text):
            if previous_was_blank or paragraph is doc.paragraphs[0]:
                paragraph._p.getparent().remove(paragraph._p)
                removed_blank_count += 1
                continue
            previous_was_blank = True
            continue
        previous_was_blank = False

        if style_name in HEADING_COLORS or style_name in LIST_LEVELS:
            paragraph.clear()
            paragraph.add_run(clean_text(text))
            continue

        if BULLET_RE.match(text):
            make_list(paragraph, BULLET_RE, "List Bullet")
            continue

        if NUMBER_RE.match(text):
            make_list(paragraph, NUMBER_RE, "List Number")
            continue

        paragraph.clear()
        paragraph.add_run(clean_text(text))
        set_alignment(paragraph, style_name)

    format_tables(doc)

    doc.save(output_path)
    print(f"Done: {input_path} -> {output_path}")
    print(f"Body font: {body_name} {body_size.pt:.1f}pt")
    print(f"Removed {removed_blank_count} extra blank paragraph(s).")
    print(f"Centered {image_count} image paragraph(s).")
    print(f"Formatted {len(doc.tables)} table(s).")


def _default_output(input_path):
    if input_path.lower().endswith(".docx"):
        return input_path[:-5] + "_formatted.docx"
    return input_path + "_formatted.docx"


def main(argv):
    if len(argv) < 2 or len(argv) > 3:
        print("Usage: python format_doc.py <input.docx> [output.docx]")
        return 1

    input_path = argv[1]
    output_path = argv[2] if len(argv) == 3 else _default_output(input_path)
    try:
        format_document(input_path, output_path)
    except FileNotFoundError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 2
    except Exception as exc:  # noqa: BLE001
        print(f"Unexpected error: {exc}", file=sys.stderr)
        return 3
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
