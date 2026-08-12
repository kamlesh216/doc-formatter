"""
validator.py — Post-formatting validation.

Reopens the saved DOCX and verifies structural integrity.
"""
from __future__ import annotations

import logging
from docx import Document

log = logging.getLogger(__name__)


def validate(
    output_path: str,
    original_table_count: int,
    original_para_count: int,
    original_image_count: int,
) -> list[str]:
    """
    Reopen the output DOCX and check for suspicious changes.
    Returns a list of warning strings (empty = all ok).
    """
    warnings = []

    try:
        doc = Document(output_path)
    except Exception as exc:
        warnings.append(f"CRITICAL: Cannot reopen output file: {exc}")
        return warnings

    # 1. Table count
    new_table_count = len(doc.tables)
    if new_table_count != original_table_count:
        warnings.append(
            f"Table count changed: {original_table_count} → {new_table_count}"
        )

    # 2. Paragraph count (allow ±10% variance from blank-para removal)
    new_para_count = len(doc.paragraphs)
    if original_para_count > 0:
        ratio = new_para_count / original_para_count
        if ratio < 0.80:
            warnings.append(
                f"Paragraph count dropped significantly: "
                f"{original_para_count} → {new_para_count}"
            )

    # 3. Sections present
    if not doc.sections:
        warnings.append("No sections found in output document!")

    # 4. Image count via XML
    from docx.oxml.ns import qn
    new_image_count = sum(
        len(para._p.findall(".//" + qn("w:drawing")))
        for para in doc.paragraphs
    )
    if new_image_count < original_image_count:
        warnings.append(
            f"Image count dropped: {original_image_count} → {new_image_count}"
        )

    # 5. Not accidentally empty
    has_content = any(p.text.strip() for p in doc.paragraphs)
    if not has_content:
        warnings.append("Output document appears to have no text content!")

    # 6. Headers / footers
    for i, section in enumerate(doc.sections):
        try:
            _ = section.header
            _ = section.footer
        except Exception as exc:
            warnings.append(f"Section {i} header/footer error: {exc}")

    if warnings:
        log.warning("Validation warnings:\n  " + "\n  ".join(warnings))
    else:
        log.info("Validation passed — no issues detected.")

    return warnings


def print_validation(warnings: list[str]) -> None:
    print("\n" + "=" * 60)
    print("VALIDATION")
    print("=" * 60)
    if not warnings:
        print("  ✓ All checks passed.")
    else:
        for w in warnings:
            print(f"  ⚠  {w}")
    print("=" * 60 + "\n")
