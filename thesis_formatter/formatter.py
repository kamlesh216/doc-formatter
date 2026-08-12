"""
formatter.py — Main orchestrator.

Runs all formatting modules in the correct order:
  1. Page layout
  2. Analysis pass
  3. Headings
  4. Captions
  5. Images
  6. Tables
  7. Paragraphs (body + bibliography)
  8. Lists
  9. Equations
 10. Blank paragraph cleanup
"""
from __future__ import annotations

import logging
from dataclasses import dataclass
from docx import Document

from . import config
from .analyzer import DocAnalysis, ParaType, analyse, print_analysis
from .captions import format_captions
from .equations import format_equations
from .headings import format_headings
from .images import format_images
from .lists import format_lists
from .page_layout import apply_page_layout
from .paragraphs import format_paragraphs
from .tables import format_tables
from .utils import is_blank_para

log = logging.getLogger(__name__)


@dataclass
class FormatReport:
    input_path:        str = ""
    output_path:       str = ""
    para_count:        int = 0
    tables_formatted:  int = 0
    images_centered:   int = 0
    images_resized:    int = 0
    figures_detected:  int = 0
    captions_formatted:int = 0
    headings_formatted:int = 0
    blanks_removed:    int = 0
    
    sections_processed:int = 0
    analysis: DocAnalysis | None = None

    def print(self) -> None:
        print("\n" + "=" * 60)
        print("FORMATTING REPORT")
        print("=" * 60)
        print(f"  Input           : {self.input_path}")
        print(f"  Output          : {self.output_path}")
        print(f"  Sections        : {self.sections_processed}")
        print(f"  Paragraphs      : {self.para_count}")
        print(f"  Tables formatted: {self.tables_formatted}")
        print(f"  Images centered : {self.images_centered}")
        print(f"  Images resized  : {self.images_resized}")
        print(f"  Figures detected: {self.figures_detected}")
        print(f"  Captions fmt    : {self.captions_formatted}")
        print(f"  Headings fmt    : {self.headings_formatted}")
        print(f"  Blanks removed  : {self.blanks_removed}")
        print("=" * 60 + "\n")


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def format_document(
    input_path: str,
    output_path: str,
    show_analysis: bool = False,
) -> FormatReport:
    """
    Load input_path, apply all formatting, save to output_path.
    Returns a FormatReport.
    """
    import os
    if not os.path.isfile(input_path):
        raise FileNotFoundError(f"Input file not found: {input_path}")

    log.info("Loading: %s", input_path)
    doc = Document(input_path)

    report = FormatReport(input_path=input_path, output_path=output_path)

    # --- Step 1: Analyse ------------------------------------------------
    log.info("Analysing document…")
    analysis = analyse(doc)
    report.analysis    = analysis
    report.para_count  = analysis.para_count

    if show_analysis:
        print_analysis(analysis)

    # --- Step 2: Page layout -------------------------------------------
    report.sections_processed = apply_page_layout(doc)

    # --- Step 3: Headings ----------------------------------------------
    report.headings_formatted = format_headings(doc, analysis)

    # --- Step 4: Captions ----------------------------------------------
    report.captions_formatted = format_captions(doc, analysis)

    # --- Step 5: Images ------------------------------------------------
    centered, resized = format_images(doc, analysis)
    report.images_centered = centered
    report.images_resized  = resized
    report.figures_detected = analysis.figure_cap_count

    # --- Step 6: Tables ------------------------------------------------
    report.tables_formatted = format_tables(doc)

    # --- Step 7: Body paragraphs + bibliography ------------------------
    format_paragraphs(doc, analysis)

    # --- Step 8: Lists -------------------------------------------------
    format_lists(doc, analysis)

    # --- Step 9: Equations ---------------------------------------------
    format_equations(doc, analysis)

    # --- Step 10: Blank paragraph cleanup ------------------------------
    report.blanks_removed = _remove_excess_blanks(doc)

    # --- Save ----------------------------------------------------------
    log.info("Saving: %s", output_path)
    doc.save(output_path)

    return report


# ---------------------------------------------------------------------------
# Blank paragraph cleanup
# ---------------------------------------------------------------------------

def _remove_excess_blanks(doc: Document) -> int:
    """
    Remove consecutive blank paragraphs beyond MAX_CONSECUTIVE_BLANK_PARAS.
    Returns count removed.
    """
    max_consecutive = config.MAX_CONSECUTIVE_BLANK_PARAS
    removed  = 0
    streak   = 0
    to_remove = []

    for para in doc.paragraphs:
        if is_blank_para(para):
            streak += 1
            if streak > max_consecutive:
                to_remove.append(para)
        else:
            streak = 0

    for para in to_remove:
        parent = para._p.getparent()
        if parent is not None:
            parent.remove(para._p)
            removed += 1

    log.info("Removed %d excess blank paragraph(s).", removed)
    return removed
